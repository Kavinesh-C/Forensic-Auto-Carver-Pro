import os
import zipfile
import mimetypes
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from flask import Flask, render_template_string, request, redirect, url_for, flash, jsonify, Response, send_from_directory, make_response, send_file
from werkzeug.utils import secure_filename
from PIL import Image, ImageTk
import io
import threading
import string
import pytsk3
import hashlib
import binascii
import math
import struct
import magic
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
import base64
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.backends import default_backend
import zlib
import re
import datetime
# --- NEW IMPORTS ---
import json
import csv
import xml.etree.ElementTree as ET
import mmap  # PERFORMANCE: Import mmap for efficient file handling

# Add with your other imports
try:
    from weasyprint import HTML
except ImportError:
    HTML = None # Handle optional dependency

try:
    import docx
except ImportError:
    docx = None # Handle optional dependency

# Optional: For Windows Event Logs (.evtx)
try:
    from Evtx.Evtx import Evtx
except ImportError:
    Evtx = None
    # Optionally, print a warning for the developer
    print("Warning: 'python-evtx' is not installed. .evtx log parsing will not be available. Install with: pip install python-evtx")
# --- END NEW IMPORTS ---


# --- Application Setup ---
# Initialize the Flask application
app = Flask(__name__)
# Secret key for session management (e.g., flash messages)
app.secret_key = 'supersecretkey'
# Configuration for file uploads
UPLOAD_FOLDER = 'uploads'
CARVED_FOLDER = 'carved_files'
DECRYPTED_FOLDER = 'decrypted_files'
DICTIONARY_FILE = 'common_passwords.txt'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CARVED_FOLDER'] = CARVED_FOLDER
app.config['DECRYPTED_FOLDER'] = DECRYPTED_FOLDER
app.config['DICTIONARY_FILE'] = DICTIONARY_FILE

# Ensure directories exist when the application starts
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CARVED_FOLDER'], exist_ok=True)
os.makedirs(app.config['DECRYPTED_FOLDER'], exist_ok=True)

# Create a default dictionary file if it doesn't exist
if not os.path.exists(DICTIONARY_FILE):
    with open(DICTIONARY_FILE, 'w') as f:
        f.write("password\n123456\nadmin\n12345\n12345678\nletmein\nqwerty\npassword1\n")

# --- In-Memory Database (for demonstration) ---
uploaded_files_db = {}
carved_files_db = {}
deleted_files_db = {} # Database for pytsk3 deleted files

# --- Status Dictionaries ---
carving_status = {
    "progress": 0, "current_offset": "0x00000000", "files_found": 0,
    "complete": False, "found_files_list": []
}
deleted_scan_status = {
    "in_progress": False, "files_found": 0, "complete": False, "message": "Scan has not started."
}
decryption_status = {
    "in_progress": False, "complete": False, "message": "", "attempts": 0
}

# --- Encryption Detection Patterns ---
ENCRYPTION_SIGNATURES = {
    'AES_ENCRYPTED': {'header': b'Salted__', 'description': 'OpenSSL AES encrypted file'},
    'BITLOCKER': {'header': b'-FVE-FS-', 'description': 'BitLocker encrypted volume'},
    'VERACRYPT': {'header': b'VERA', 'description': 'VeraCrypt encrypted container'},
    'PGP': {'header': b'\x85\x01\x0c', 'description': 'PGP encrypted file'},
    '7Z_ENCRYPTED': {'header': b'7z\xbc\xaf\x27\x1c', 'description': '7-Zip encrypted archive'},
    'ZIP_ENCRYPTED': {'header': b'PK\x03\x04', 'description': 'ZIP encrypted archive (needs password)'},
    'RAR_ENCRYPTED': {'header': b'Rar!\x1a\x07', 'description': 'RAR encrypted archive'},
}

# Common passwords for dictionary attacks
COMMON_PASSWORDS = [
    "password", "123456", "12345678", "qwerty", "abc123", "monkey", "letmein",
    "dragon", "111111", "baseball", "iloveyou", "trustno1", "sunshine", "master",
    "123123", "welcome", "shadow", "ashley", "football", "jesus", "michael", "ninja",
    "mustang", "password1", "admin", "1234", "12345", "123456789", "1234567",
]

# --- MODIFIED: File Signature Database with Log Files ---
FILE_SIGNATURES = {
    'Image Files': {
        'JPG': {'header': b'\xff\xd8\xff', 'footer': b'\xff\xd9', 'extension': '.jpg'},
        'JPEG': {'header': b'\xff\xd8\xff', 'footer': b'\xff\xd9', 'extension': '.jpeg'},
        'PNG': {'header': b'\x89\x50\x4e\x47\x0d\x0a\x1a\x0a', 'footer': b'\x49\x45\x4e\x44\xae\x42\x60\x82', 'extension': '.png'},
        'GIF': {'header': b'\x47\x49\x46\x38', 'footer': b'\x00\x3b', 'extension': '.gif'},
        'TIFF (Intel)': {'header': b'\x49\x49\x2a\x00', 'footer': None, 'max_size': 30*1024*1024, 'extension': '.tiff'},
        'TIFF (Motorola)': {'header': b'\x4d\x4d\x00\x2a', 'footer': None, 'max_size': 30*1024*1024, 'extension': '.tiff'},
        'BMP': {'header': b'\x42\x4d', 'footer': None, 'max_size': 30*1024*1024, 'extension': '.bmp'},
        'ICO': {'header': b'\x00\x00\x01\x00', 'footer': None, 'max_size': 1*1024*1024, 'extension': '.ico'},
    },
    'Document Files': {
        'PDF': {'header': b'\x25\x50\x44\x46', 'footer': b'\x25\x25\x45\x4f\x46', 'extension': '.pdf'},
        'RTF': {'header': b'\x7b\x5c\x72\x74\x66', 'footer': b'\x7d', 'extension': '.rtf'},
        'DOC': {'header': b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1', 'footer': None, 'max_size': 20*1024*1024, 'extension': '.doc'},
        'UTF-8 BOM Text': {'header': b'\xef\xbb\xbf', 'footer': None, 'max_size': 5*1024*1024, 'extension': '.txt'},
        'UTF-16LE BOM Text': {'header': b'\xff\xfe', 'footer': None, 'max_size': 5*1024*1024, 'extension': '.txt'},
        'UTF-16BE BOM Text': {'header': b'\xfe\xff', 'footer': None, 'max_size': 5*1024*1024, 'extension': '.txt'},
        'UTF-32LE BOM Text': {'header': b'\xff\xfe\x00\x00', 'footer': None, 'max_size': 5*1024*1024, 'extension': '.txt'},
        'UTF-32BE BOM Text': {'header': b'\x00\x00\xfe\xff', 'footer': None, 'max_size': 5*1024*1024, 'extension': '.txt'},
    },
    'Archive Files': {
        'ZIP': {'header': b'\x50\x4b\x03\x04', 'footer': b'\x50\x4b\x05\x06', 'extension': '.zip'},
        'RAR': {'header': b'\x52\x61\x72\x21\x1a\x07\x00', 'footer': None, 'max_size': 1024*1024*1024, 'extension': '.rar'},
        '7Z': {'header': b'\x37\x7a\xbc\xaf\x27\x1c', 'footer': None, 'max_size': 1024*1024*1024, 'extension': '.7z'},
        'GZIP': {'header': b'\x1f\x8b', 'footer': None, 'max_size': 1024*1024*1024, 'extension': '.gz'},
    },
    'Email Files': {
        'Outlook Data File (PST/OST)': {'header': b'!\x42\x44\x4e', 'footer': None, 'max_size': 20 * 1024 * 1024 * 1024, 'extension': '.pst'},
        'MBOX Email Archive': {'header': b'From ', 'footer': None, 'max_size': 10 * 1024 * 1024 * 1024, 'extension': '.mbox'},
    },
    'Network Files': {
        'PCAP Capture File': {'header': b'\xd4\xc3\xb2\xa1', 'footer': None, 'max_size': 2048 * 1024 * 1024, 'extension': '.pcap'},
        'PCAPng Capture File': {'header': b'\x0a\x0d\x0d\x0a', 'footer': None, 'max_size': 2048 * 1024 * 1024, 'extension': '.pcapng'},
    },
    'Audio Files': {
        'MP3 (ID3)': {'header': b'\x49\x44\x33', 'footer': None, 'max_size': 15*1024*1024, 'extension': '.mp3'},
        'WAV': {'header': b'\x52\x49\x46\x46', 'footer': None, 'max_size': 100*1024*1024, 'extension': '.wav'},
        'FLAC': {'header': b'\x66\x4c\x61\x43', 'footer': None, 'max_size': 100*1024*1024, 'extension': '.flac'},
    },
    'Video Files': {
        'MP4': {'header': b'\x00\x00\x00\x18\x66\x74\x79\x70\x69\x73\x6f\x6d', 'footer': None, 'max_size': 1024*1024*1024, 'extension': '.mp4'},
        'AVI': {'header': b'\x52\x49\x46\x46', 'footer': None, 'max_size': 1024*1024*1024, 'extension': '.avi'},
        'MOV': {'header': b'\x00\x00\x00\x14\x66\x74\x79\x70\x71\x74\x20\x20', 'footer': None, 'max_size': 1024*1024*1024, 'extension': '.mov'},
    },
    'Executable Files': {
        'EXE': {'header': b'\x4d\x5a', 'footer': None, 'max_size': 50*1024*1024, 'extension': '.exe'},
        'ELF': {'header': b'\x7f\x45\x4c\x46', 'footer': None, 'max_size': 50*1024*1024, 'extension': '.elf'},
    },
    'Database Files': {
        'SQLite': {'header': b'\x53\x51\x4c\x69\x74\x65\x20\x66\x6f\x72\x6d\x61\x74\x20\x33\x00', 'footer': None, 'max_size': 100*1024*1024, 'extension': '.sqlite'},
        'MDB': {'header': b'\x00\x01\x00\x00\x53\x74\x61\x6E\x64\x61\x72\x64\x20\x4A\x65\x74\x20\x44\x42', 'footer': None, 'max_size': 500*1024*1024, 'extension': '.mdb'},
        'ACCDB': {'header': b'\x00\x01\x00\x00\x53\x74\x61\x6E\x64\x61\x72\x64\x20\xAC\x45\x44\x42', 'footer': None, 'max_size': 500*1024*1024, 'extension': '.accdb'},
    },
    'Compiled/Binary Files': {
        'Java Class': {'header': b'\xca\xfe\xba\xbe', 'footer': None, 'max_size': 10*1024*1024, 'extension': '.class'},
    },

    'Windows Log Files': {
        'Windows Event Log (EVTX)': {
            'header': b'\x45\x6c\x66\x46\x69\x6c\x65\x00',
            'size_from_header': {'offset': 8, 'format': '<Q'}, # More accurate carving
            'extension': '.evtx'
        },
        'Windows Event Log (EVT)': {
            'header': b'\x4c\x66\x4c\x65\x00\x00\x00\x00',
            'size_from_header': {'offset': 4, 'format': '<L'}, # Accurate carving for legacy format
            'extension': '.evt'
        },
        'Windows Registry Hive (REGF)': {
            'header': b'\x72\x65\x67\x66',
            'max_size': 500 * 1024 * 1024, # Hives can be large, size is complex to parse
            'extension': '.dat'
        },
        'Windows Prefetch (PF)': {
            'header': b'SCCA',
            'size_from_header': {'offset': 12, 'format': '<L'},
            'extension': '.pf'
        },
        'Windows IIS Log': {
            'header': b'#Software: Microsoft Internet Information Services',
            'max_size': 100 * 1024 * 1024, # Text-based, so use max size
            'extension': '.log'
        },
        'Windows Event Log (EVTX)': {
            'header': b'\x45\x6c\x66\x46\x69\x6c\x65\x00',
            'size_from_header': {'offset': 8, 'format': '<Q'}, # More accurate carving
            'extension': '.evtx'
        },
        'Windows Event Log (EVT)': {
            'header': b'\x4c\x66\x4c\x65\x00\x00\x00\x00',
            'size_from_header': {'offset': 4, 'format': '<L'}, # Accurate carving for legacy format
            'extension': '.evt'
        },
    },
    'Linux Log Files': {
        'Linux Syslog': {
            'header': b'<134>',
            'max_size': 100 * 1024 * 1024, # Text-based, so use max size
            'extension': '.log'
        },
        'Linux Journal': {
            'header': b'Journal',
            'max_size': 100 * 1024 * 1024, # Text-based, so use max size
            'extension': '.log'
        },
        'Linux Audit Log': {
            'header': b'type=audit',
            'max_size': 100 * 1024 * 1024, # Text-based, so use max size
            'extension': '.log'
        },
        },
        'Mac os Log Files': {
            'header': b'OS X ',
            'max_size': 100 * 1024 * 1024, # Text-based, so use max size
            'extension': '.log'
        },
    
        # NOTE: Generic text logs like Linux syslog or macOS system.log have no standard
        # header and cannot be reliably carved with this signature-based method.
        # Use the BOM Text/Log option below for compatible text files.

}


LOG_OS_OPTIONS = {
    "Windows": [".log", ".evtx", ".txt", ".csv", ".json", ".xml"],
    "Linux": [".log", ".txt", ".json", ".xml", ".csv", ".gz", ".zip", ""],
    "macOS": [".logarchive", ".gz", ".tracev3", ".log"]
}
EVENT_OS_OPTIONS = {
    "Windows": [".evtx", ".evt"],
    "Linux": ["/var/log/syslog", "/var/log/auth.log", "/var/log/kern.log",
              "/var/log/dmesg", "/var/log/wtmp", "/var/log/btmp"],
    "macOS": [".logarchive", ".asl", ".ips", ".spin", ".diag", ".log"]
}

# --- NEW HELPER FUNCTIONS FOR LOG PARSING ---
def parse_text_log(filepath):
    with open(filepath, "r", errors="ignore") as f:
        data = f.read()
    return data

def create_thumbnail_data_uri(filepath):
    """Creates a Base64-encoded data URI for an image thumbnail."""
    try:
        with Image.open(filepath) as img:
            img.thumbnail((100, 100)) # Create a 100x100 thumbnail
            buffered = io.BytesIO()
            img.save(buffered, format="PNG")
            img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
            return f"data:image/png;base64,{img_str}"
    except (IOError, OSError):
        return None # Return None if it's not a valid image file
    

def parse_csv_log(filepath):
    output = []
    with open(filepath, newline="", errors="ignore") as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            output.append(",".join(row))
    return "\n".join(output)

def parse_json_log(filepath):
    with open(filepath, "r", errors="ignore") as f:
        data = json.load(f)
    return json.dumps(data, indent=4)

def parse_xml_log(filepath):
    output = []
    tree = ET.parse(filepath)
    root = tree.getroot()
    for elem in root.iter():
        text = elem.text.strip() if elem.text else ""
        output.append(f"<{elem.tag}> {text}")
    return "\n".join(output)

def parse_evtx_log(filepath):
    if Evtx is None:
        return "[ERROR] The 'python-evtx' library is required to parse .evtx files. Please install it using: pip install python-evtx"
    output = []
    with Evtx(filepath) as log:
        for record in log.records():
            output.append(record.xml())
    return "\n\n".join(output)

# --- Forensic & Encryption Functions ---

def calculate_entropy(data):
    """Calculate Shannon entropy of data accurately."""
    if not data: return 0
    byte_counts = [0] * 256
    for byte in data: byte_counts[byte] += 1
    entropy = 0
    total_bytes = len(data)
    for count in byte_counts:
        if count > 0:
            probability = count / total_bytes
            entropy -= probability * math.log2(probability)
    return entropy

def detect_encryption(filepath):
    """Detects if a file is encrypted by checking signatures and entropy."""
    try:
        with open(filepath, 'rb') as f: header = f.read(512)
        for enc_type, sig_info in ENCRYPTION_SIGNATURES.items():
            if sig_info['header'] and header.startswith(sig_info['header']):
                if enc_type == 'ZIP_ENCRYPTED':
                    try:
                        with zipfile.ZipFile(filepath) as zf:
                            if not any(e.flag_bits & 0x1 for e in zf.infolist()):
                                continue
                    except Exception:
                        continue
                return {'encrypted': True, 'encryption_type': enc_type, 'description': sig_info['description']}
        entropy = calculate_entropy(header)
        if entropy > 7.5:
            return {'encrypted': True, 'encryption_type': 'UNKNOWN', 'description': f'High entropy ({entropy:.2f}) suggests encryption'}
    except Exception as e:
        print(f"Error in detect_encryption: {e}")
    return {'encrypted': False, 'encryption_type': None, 'description': 'No encryption detected'}

def perform_forensic_analysis(file_path):
    """Perform basic forensic analysis on an image."""
    results = []
    try:
        with open(file_path, 'rb') as f: header = f.read(4096)
        if header[510:512] == b'\x55\xaa': results.append("✓ MBR Partition Table detected")
        if b'EFI PART' in header[512:1024]: results.append("✓ GPT Partition Table detected")
        if b'FAT32' in header[0x52:0x5A]: results.append("✓ FAT32 File System detected")
        elif b'NTFS' in header[0x03:0x07]: results.append("✓ NTFS File System detected")
        elif b'\x53\xEF' in header[1024+56:1024+58]: results.append("✓ EXT File System Superblock detected")
        if b'-FVE-FS-' in header: results.append("⚠ BitLocker encryption detected")
        file_size = os.path.getsize(file_path)
        results.append(f"📊 File Size: {file_size/(1024*1024*1024):.2f} GB" if file_size > 1024**3 else f"📊 File Size: {file_size/(1024*1024):.2f} MB")
        entropy = calculate_entropy(header)
        results.append(f"📈 Entropy: {entropy:.2f} (High entropy > 7.5 may suggest encryption)")
    except Exception as e:
        results.append(f"❌ Analysis error: {str(e)}")
    return results

def show_hash_info(file_path):
    """Calculate and display file hashes."""
    hashes = {}
    try:
        with open(file_path, 'rb') as f: content = f.read(65536)
        hashes['MD5'] = hashlib.md5(content).hexdigest()
        hashes['SHA-1'] = hashlib.sha1(content).hexdigest()
        hashes['SHA-256'] = hashlib.sha256(content).hexdigest()
    except Exception as e:
        print(f"Error calculating hashes: {e}")
    return hashes

def format_hex_view(data_bytes, start_offset=0):
    """Formats a byte string into a complete hex view string."""
    lines = []
    for i in range(0, len(data_bytes), 16):
        chunk = data_bytes[i:i+16]
        offset_str = f"0x{(start_offset + i):08X}"
        hex_str = ' '.join(f'{b:02X}' for b in chunk)
        ascii_str = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in chunk)
        lines.append(f"{offset_str}  {hex_str:<47} {ascii_str}")
    return '\n'.join(lines)

def attempt_decryption(filepath, encryption_type, password=None):
    """Master decryption function that handles password and passwordless attempts."""
    global decryption_status, uploaded_files_db
    filename = os.path.basename(filepath)
    decryption_status.update({"in_progress": True, "complete": False, "message": "Starting decryption..."})
    if filename in uploaded_files_db: uploaded_files_db[filename]['encryption_status']['decrypting'] = True
    
    decrypted_path = os.path.join(app.config['DECRYPTED_FOLDER'], f"decrypted_{filename}")
    passwords_to_try = [password] if password else COMMON_PASSWORDS
    
    success = False
    if "ZIP" in encryption_type:
        success = crack_zip_with_passwords(filepath, decrypted_path, passwords_to_try, bool(password))
    elif "AES" in encryption_type:
        success = crack_openssl_aes_with_passwords(filepath, decrypted_path, passwords_to_try, bool(password))
    else:
        decryption_status['message'] = f"Automated decryption for '{encryption_type}' is not supported."
        
    decryption_status.update({'in_progress': False, 'complete': True})
    
    if success:
        decryption_status['message'] = "Decryption successful!"
        if filename in uploaded_files_db:
            uploaded_files_db[filename]['encryption_status']['decrypted_path'] = decrypted_path
            uploaded_files_db[filename]['path'] = decrypted_path
    else:
        decryption_status['message'] = "Decryption failed. All password attempts were incorrect."
    
    if filename in uploaded_files_db:
        uploaded_files_db[filename]['encryption_status']['decrypting'] = False

def crack_zip_with_passwords(filepath, output_path, passwords_to_try, is_user_pwd):
    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            for i, pwd in enumerate(passwords_to_try):
                decryption_status['attempts'] = i + 1
                decryption_status["message"] = f"Trying password: {'●' * len(pwd) if is_user_pwd else pwd}"
                try:
                    zf.extractall(os.path.dirname(output_path), pwd=pwd.encode())
                    extracted_name = zf.namelist()[0]
                    os.rename(os.path.join(os.path.dirname(output_path), extracted_name), output_path)
                    return True
                except (RuntimeError, zipfile.BadZipFile): continue
    except Exception as e: print(f"ZIP cracking error: {e}")
    return False

def crack_openssl_aes_with_passwords(filepath, output_path, passwords_to_try, is_user_pwd):
    try:
        with open(filepath, 'rb') as f: data = f.read()
        if data.startswith(b'Salted__'):
            salt, encrypted_data = data[8:16], data[16:]
            for i, pwd in enumerate(passwords_to_try):
                decryption_status.update({"attempts": i + 1, "message": f"Trying AES password: {'●' * len(pwd) if is_user_pwd else pwd}"})
                try:
                    kdf = PBKDF2HMAC(algorithm=hashes.SHA256(), length=32, salt=salt, iterations=100000)
                    key = base64.urlsafe_b64encode(kdf.derive(pwd.encode()))
                    decrypted_data = Fernet(key).decrypt(encrypted_data)
                    with open(output_path, 'wb') as f_out: f_out.write(decrypted_data)
                    return True
                except Exception: continue
    except Exception as e: print(f"AES cracking error: {e}")
    return False

# --- PERFORMANCE OPTIMIZATION using mmap ---
def simple_file_carver(filepath, selected_types):
    """File carving engine optimized with mmap for speed and memory efficiency."""
    global carving_status, carved_files_db
    carving_status = {"progress": 0, "current_offset": "0x00000000", "files_found": 0, "complete": False, "found_files_list": []}
    carved_files_db.clear()
    found_file_counter = 0

    signatures_to_check = [(name, sig) for cat in FILE_SIGNATURES.values() for name, sig in cat.items() if name in selected_types and sig.get('header')]

    try:
        with open(filepath, 'rb') as f:
            with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                file_size = len(mm)
                current_pos = 0

                while current_pos < file_size:
                    carving_status["progress"] = int((current_pos / file_size) * 100) if file_size > 0 else 100
                    carving_status["current_offset"] = f"0x{current_pos:08X}"

                    # Find the next potential header from the current position
                    found_positions = []
                    for name, sig in signatures_to_check:
                        header_pos = mm.find(sig['header'], current_pos)
                        if header_pos != -1:
                            found_positions.append({'pos': header_pos, 'name': name, 'sig': sig})
                    
                    if not found_positions:
                        break # No more headers found

                    # Get the earliest found header
                    next_file = min(found_positions, key=lambda x: x['pos'])
                    start_pos = next_file['pos']
                    name, sig = next_file['name'], next_file['sig']
                    
                    end_pos = -1
                    if sig.get('footer'):
                        footer_pos = mm.find(sig['footer'], start_pos + len(sig['header']))
                        if footer_pos != -1:
                            end_pos = footer_pos + len(sig['footer'])
                    elif sig.get('max_size'):
                        # For size-limited carving, find the start of the *next* known file type to avoid overlap
                        next_header_starts = []
                        for other_name, other_sig in signatures_to_check:
                           next_h_pos = mm.find(other_sig['header'], start_pos + len(sig['header']))
                           if next_h_pos != -1:
                               next_header_starts.append(next_h_pos)
                        
                        max_end = start_pos + sig['max_size']
                        if next_header_starts:
                            # End before the earliest next header, if it's within the max_size
                            end_pos = min(min(next_header_starts), max_end)
                        else:
                            end_pos = max_end
                    
                    if end_pos != -1:
                        end_pos = min(end_pos, file_size) # Ensure we don't read past EOF
                        found_file_counter += 1
                        content = mm[start_pos:end_pos]
                        filename = f"carved_{found_file_counter}{sig['extension']}"
                        save_path = os.path.join(app.config['CARVED_FOLDER'], filename)
                        with open(save_path, 'wb') as carved_file:
                            carved_file.write(content)
                        
                        file_info = {"id": found_file_counter, "name": filename, "offset": f"0x{start_pos:08X}", "size_kb": f"{len(content)/1024:.2f} KB", "hex_preview": format_hex_view(content[:256])}
                        carved_files_db[filename] = file_info
                        carving_status["found_files_list"].append(file_info)
                        carving_status["files_found"] = found_file_counter
                        
                        current_pos = end_pos
                    else:
                        # If a header was found but no valid end, advance past the header to avoid an infinite loop
                        current_pos = start_pos + 1
                        
    except Exception as e:
        print(f"Error during carving: {e}")
    
    carving_status.update({"progress": 100, "complete": True})

def scan_for_deleted_files_engine(filepath):
    """Scans a disk image for deleted files using pytsk3 and updates global state."""
    global deleted_files_db, deleted_scan_status
    deleted_scan_status.update({"in_progress": True, "complete": False, "files_found": 0, "message": "Starting scan..."})
    found_files = {}
    try:
        img_handle = pytsk3.Img_Info(filepath)
        def walk_and_save(fs, directory, fs_offset):
            for f in directory:
                if not hasattr(f.info, 'meta') or f.info.meta is None: continue
                is_deleted = not (f.info.meta.flags & pytsk3.TSK_FS_META_FLAG_ALLOC)
                is_file = f.info.meta.type == pytsk3.TSK_FS_META_TYPE_REG
                has_name = hasattr(f.info, 'name') and f.info.name is not None
                if is_deleted and is_file and f.info.meta.size > 0:
                    name = f.info.name.name.decode('utf-8', 'ignore') if has_name else '[orphan_file]'
                    inode = f.info.meta.addr
                    found_files[inode] = {
                      "name": name, "size": f.info.meta.size, "inode": inode, "fs_offset": fs_offset,
                      "mtime": datetime.datetime.fromtimestamp(f.info.meta.mtime).strftime('%Y-%m-%d %H:%M:%S'),
                        "atime": datetime.datetime.fromtimestamp(f.info.meta.atime).strftime('%Y-%m-%d %H:%M:%S'),
                        "ctime": datetime.datetime.fromtimestamp(f.info.meta.ctime).strftime('%Y-%m-%d %H:%M:%S')
                    }
                    deleted_scan_status["files_found"] = len(found_files)
                    deleted_scan_status["message"] = f"Found {len(found_files)} deleted file entries..."
                if f.info.meta.type == pytsk3.TSK_FS_META_TYPE_DIR and has_name and f.info.name.name not in [b'.', b'..']:
                    try: walk_and_save(fs, f.as_directory(), fs_offset)
                    except IOError: pass
        try:
            volume = pytsk3.Volume_Info(img_handle)
            for part in volume:
                if part.flags != pytsk3.TSK_VS_PART_FLAG_UNALLOC:
                    try:
                        fs_offset = part.start * volume.info.block_size
                        fs = pytsk3.FS_Info(img_handle, offset=fs_offset)
                        walk_and_save(fs, fs.open_dir(path="/"), fs_offset)
                    except IOError: continue
        except IOError:
            try:
                fs = pytsk3.FS_Info(img_handle, offset=0)
                walk_and_save(fs, fs.open_dir(path="/"), 0)
            except IOError as e: deleted_scan_status["message"] = f"Error opening as filesystem: {e}. The disk image might be encrypted or unsupported."
        deleted_files_db = found_files
        deleted_scan_status["message"] = f"Scan complete. Found {len(found_files)} deleted file entries."
    except Exception as e:
        print(f"Error scanning for deleted files: {e}")
        deleted_scan_status["message"] = f"An error occurred: {e}"
    deleted_scan_status.update({"in_progress": False, "complete": True})

def generate_docx_report_data(case_details, evidence_file, carved_files, deleted_files, now):
    """Generates a .docx report using the python-docx library."""
    if not docx:
        return None # Return None if the library is not installed

    document = docx.Document()
    document.add_heading('Forensic Report', 0)
    document.add_paragraph(f"Generated by ForensicCarver Pro on {now.strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_heading('Case Details', level=1)
    document.add_paragraph(f"Case Name: {case_details['name']}")
    document.add_paragraph(f"Case Number: {case_details['number']}")
    document.add_paragraph(f"Examiner: {case_details['examiner']}")

    for filename, details in evidence_file.items():
        document.add_heading(f"Data Source: {filename}", level=1)
        document.add_paragraph(f"File Size: {details['size_mb']} MB")
        document.add_paragraph(f"MD5 Hash: {details['hash_info']['MD5']}")
        document.add_paragraph(f"SHA-1 Hash: {details['hash_info']['SHA-1']}")
        
        document.add_heading('Partition Information', level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Index'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Start Sector'
        hdr_cells[3].text = 'Length (Sectors)'
        for part in details['partition_info']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(part.get('addr', ''))
            row_cells[1].text = part.get('desc', '')
            row_cells[2].text = str(part.get('start', ''))
            row_cells[3].text = str(part.get('len', ''))

    document.add_heading(f"Carved Files ({len(carved_files)})", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Filename'
    hdr_cells[2].text = 'Offset'
    for file in sorted(carved_files.values(), key=lambda x: x.get('id', 0)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(file.get('id', ''))
        row_cells[1].text = file.get('name', '')
        row_cells[2].text = file.get('offset', '')

    document.add_heading(f"Deleted File Entries ({len(deleted_files)})", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Inode'
    hdr_cells[1].text = 'Filename'
    hdr_cells[2].text = 'Size (Bytes)'
    hdr_cells[3].text = 'Modified'
    hdr_cells[4].text = 'Accessed'
    hdr_cells[5].text = 'Created'
    for file in deleted_files.values():
        row_cells = table.add_row().cells
        row_cells[0].text = str(file.get('inode', ''))
        row_cells[1].text = file.get('name', '')
        row_cells[2].text = str(file.get('size', ''))
        row_cells[3].text = file.get('mtime', '')
        row_cells[4].text = file.get('atime', '')
        row_cells[5].text = file.get('ctime', '')

    # Save document to a memory buffer
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --- MODIFIED: HTML Templates with Reordered Sidebar ---

BASE_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>ForensicCarver Pro</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style> body { background-color: #111827; color: #d1d5db; font-family: 'Inter', sans-serif; } .sidebar { background-color: #1f2937; } .sidebar a { border-left: 3px solid transparent; transition: all 0.2s ease-in-out; } .sidebar a.active { border-left-color: #3b82f6; background-color: #374151; color: white; } .card { background-color: #1f2937; border: 1px solid #374151; box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1), 0 2px 4px -2px rgb(0 0 0 / 0.1); } .btn-primary, .btn-secondary, .btn-green, .bg-red-600 { transition: all 0.2s ease-in-out; } .btn-primary:hover, .btn-secondary:hover, .btn-green:hover, .bg-red-600:hover { opacity: 0.9; transform: translateY(-1px); } .btn-primary { background-color: #3b82f6; color: white; } .btn-secondary { background-color: #4b5563; color: white; } .btn-green { background-color: #10b981; color: white; } .hex-view { font-family: 'Courier New', Courier, monospace; background-color: #0d1117; border: 1px solid #374151; } .log-view { font-family: 'Courier New', Courier, monospace; background-color: #0d1117; color: #d1d5db; border: 1px solid #374151; white-space: pre-wrap; word-wrap: break-word; } .encryption-badge { background-color: #ef4444; color: white; padding: 2px 8px; border-radius: 4px; font-size: 0.75rem; font-weight: bold; } .decryption-badge { background-color: #10b981; color: white; padding: 2px 8px; border-radius: 4px; font-size: 0.75rem; font-weight: bold; } .decrypting-badge { background-color: #f59e0b; color: white; padding: 2px 8px; border-radius: 4px; font-size: 0.75rem; font-weight: bold; } .modal { display: none; position: fixed; z-index: 1000; left: 0; top: 0; width: 100%; height: 100%; background-color: rgba(0,0,0,0.7); } .modal-content { background-color: #1f2937; margin: 15% auto; padding: 20px; border: 1px solid #374151; width: 50%; border-radius: 8px; box-shadow: 0 4px 8px 0 rgba(0,0,0,0.2); } .close { color: #aaa; float: right; font-size: 28px; font-weight: bold; cursor: pointer; } .close:hover { color: white; } </style>
    <link rel="stylesheet" href="https://rsms.me/inter/inter.css">
</head>
<body class="flex h-screen">
    <aside class="sidebar w-64 p-4 space-y-2 flex-shrink-0">
        <div class="text-white text-2xl font-bold mb-8">ForensicCarver <span class="text-blue-500">Pro</span><p class="text-xs font-normal text-gray-400">Digital Evidence Analysis</p></div>
        <a href="{{ url_for('evidence_upload') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'evidence_upload' %}active{% endif %}">Evidence Upload</a>
        <a href="{{ url_for('forensic_analysis') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'forensic_analysis' %}active{% endif %}">Forensic Analysis</a>
        <a href="{{ url_for('auto_carving_setup') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'auto_carving_setup' %}active{% endif %}">Auto Carving</a>
        <a href="{{ url_for('recovered_files') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'recovered_files' %}active{% endif %}">Recovered Files</a>
        <a href="{{ url_for('deleted_files') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'deleted_files' %}active{% endif %}">Deleted Files</a>
        <a href="{{ url_for('log_file_viewer') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'log_file_viewer' %}active{% endif %}">Log File Viewer</a>
        <a href="{{ url_for('event_log_viewer') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'event_log_viewer' %}active{% endif %}">Event Log Viewer</a>
        <a href="{{ url_for('reporting') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'reporting' %}active{% endif %}">Reporting</a>
        <a href="{{ url_for('manual_carving') }}" class="block p-3 rounded-lg hover:bg-gray-700 {% if request.endpoint == 'manual_carving' %}active{% endif %}">Manual Carving</a>
        <div class="absolute bottom-4 text-xs text-gray-500">Version 4.2.0<br>Licensed to: Forensics Lab</div>
    </aside>
    <main class="flex-1 p-8 overflow-y-auto">
        {% with messages = get_flashed_messages(with_categories=true) %}{% if messages %}{% for category, message in messages %}<div class="bg-blue-600 text-white p-4 rounded-lg mb-4">{{ message }}</div>{% endfor %}{% endif %}{% endwith %}
        {{ content|safe }}
    </main>
    <div id="encryptionModal" class="modal"><div class="modal-content"><span class="close">&times;</span><h2 id="modalTitle" class="text-xl font-semibold text-white mb-4"></h2><div id="modalBody" class="text-gray-300 mb-6"></div></div></div>
    <script>var modal = document.getElementById("encryptionModal"); var span = document.getElementsByClassName("close")[0]; span.onclick = function() { modal.style.display = "none"; }
        window.onclick = function(event) { if (event.target == modal) { modal.style.display = "none"; } }
        function showEncryptionModal(title, bodyHtml) { document.getElementById("modalTitle").innerText = title; document.getElementById("modalBody").innerHTML = bodyHtml; modal.style.display = "block"; }
        {% if show_encryption_modal %}showEncryptionModal('{{ modal_title|safe }}', `{{ modal_body|safe }}`);{% endif %}
    </script>
</body>
</html>
"""

REPORTING_PAGE_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Generate Case Report</h1>
<p class="text-gray-400 mb-8">Fill in the case details below, select a format, and click "Generate" to create a complete report of all findings.</p>
<div class="card p-6 rounded-lg">
    <form action="{{ url_for('reporting') }}" method="post" class="space-y-4">
        <div>
            <label for="case_name" class="block text-sm font-medium text-gray-300">Case Name</label>
            <input type="text" name="case_name" id="case_name" value="Case_{{ now.strftime('%Y%m%d') }}" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
        </div>
        <div>
            <label for="case_number" class="block text-sm font-medium text-gray-300">Case Number</label>
            <input type="text" name="case_number" id="case_number" value="{{ now.strftime('%H%M%S') }}" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
        </div>
        <div>
            <label for="examiner_name" class="block text-sm font-medium text-gray-300">Examiner Name</label>
            <input type="text" name="examiner_name" id="examiner_name" placeholder="Enter your name" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
        </div>
        <div>
            <label for="report_format" class="block text-sm font-medium text-gray-300">Report Format</label>
            <select name="report_format" id="report_format" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2 h-10">
                <option value="html">HTML</option>
                <option value="pdf">PDF</option>
                <option value="docx">Word (DOCX)</option>
                <option value="csv">CSV (Zipped)</option>
            </select>
        </div>
        <div class="pt-4">
            <button type="submit" class="btn-green w-full py-3 rounded-lg font-semibold">Generate Report</button>
        </div>
    </form>
</div>
"""

REPORT_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Forensic Report: {{ case_details.name }}</title>
    <style>
        body { font-family: sans-serif; margin: 2em; background-color: #fdfdfd; color: #333; }
        h1, h2, h3 { color: #2c3e50; border-bottom: 1px solid #ccc; padding-bottom: 5px; }
        table { width: 100%; border-collapse: collapse; margin-bottom: 2em; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; }
        tr:nth-child(even) { background-color: #f9f9f9; }
        .mono { font-family: monospace; }
        .header { background-color: #34495e; color: white; padding: 1em; margin-bottom: 2em; }
        .header h1 { color: white; border: none; }
        .thumbnail { max-width: 100px; max-height: 100px; }
    </style>
</head>
<body>
    <div class="header">
        <h1>Forensic Report</h1>
        <p>Generated by ForensicCarver Pro on {{ now.strftime('%Y-%m-%d %H:%M:%S') }}</p>
    </div>

    <h2>Case Details</h2>
    <table>
        <tr><th>Case Name</th><td>{{ case_details.name }}</td></tr>
        <tr><th>Case Number</th><td>{{ case_details.number }}</td></tr>
        <tr><th>Examiner</th><td>{{ case_details.examiner }}</td></tr>
    </table>

    {% for filename, details in evidence_file.items() %}
    <h2>Data Source: {{ filename }}</h2>
    <table>
        <tr><th style="width:20%;">Property</th><th>Value</th></tr>
        <tr><td>File Size</td><td>{{ details.size_mb }} MB</td></tr>
        <tr><td class="mono">MD5 Hash</td><td class="mono">{{ details.hash_info.MD5 }}</td></tr>
        <tr><td class="mono">SHA-1 Hash</td><td class="mono">{{ details.hash_info['SHA-1'] }}</td></tr>
        <tr><td class="mono">SHA-256 Hash</td><td class="mono">{{ details.hash_info['SHA-256'] }}</td></tr>
    </table>
    
    <h3>Partition Information</h3>
    <table>
        <thead><tr><th>Index</th><th>Description</th><th>Start Sector</th><th>Length (Sectors)</th></tr></thead>
        <tbody>
        {% for part in details.partition_info %}
            <tr><td>{{ part.addr }}</td><td>{{ part.desc }}</td><td>{{ part.start }}</td><td>{{ part.len }}</td></tr>
        {% endfor %}
        </tbody>
    </table>
    {% endfor %}

    <h2>Carved Files ({{ carved_files|length }} found)</h2>
    <table>
        <thead><tr><th>ID</th><th>Filename</th><th>Offset</th><th>Preview</th></tr></thead>
        <tbody>
        {% for file in carved_files.values()|sort(attribute='id') %}
            <tr>
                <td>{{ file.id }}</td>
                <td class="mono">{{ file.name }}</td>
                <td class="mono">{{ file.offset }}</td>
                <td>
                    {% if file.thumbnail_uri %}
                        <img src="{{ file.thumbnail_uri }}" alt="thumbnail" class="thumbnail">
                    {% else %}
                        N/A
                    {% endif %}
                </td>
            </tr>
        {% else %}
            <tr><td colspan="4">No files were carved from this image.</td></tr>
        {% endfor %}
        </tbody>
    </table>

    <h2>Deleted File Entries ({{ deleted_files|length }} found)</h2>
    <table>
        <thead><tr><th>Inode</th><th>Filename</th><th>Size (Bytes)</th><th>Modified</th><th>Accessed</th><th>Created</th></tr></thead>
        <tbody>
        {% for inode, file in deleted_files.items() %}
            <tr>
                <td class="mono">{{ file.inode }}</td>
                <td class="mono">{{ file.name }}</td>
                <td>{{ file.size }}</td>
                <td class="mono">{{ file.mtime }}</td>
                <td class="mono">{{ file.atime }}</td>
                <td class="mono">{{ file.ctime }}</td>
            </tr>
        {% else %}
            <tr><td colspan="6">No deleted file entries were found in the filesystem.</td></tr>
        {% endfor %}
        </tbody>
    </table>
</body>
</html>
"""

SCAN_PROGRESS_CONTENT = """
<div class="flex justify-center items-center h-full">
    <div class="spinner-border text-green-500" role="status">
        <span class="sr-only">Scanning...</span>
    </div>
</div>
"""

EVIDENCE_UPLOAD_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Evidence File Upload</h1>
<p class="text-gray-400 mb-8">Upload digital evidence files for analysis. Supported formats: .dd, .e01, .mem, .raw, .img, .vmdk</p>
<div class="grid grid-cols-1 lg:grid-cols-2 gap-8">
    <div class="space-y-8">
        <div class="card p-6 rounded-lg">
            <h2 class="text-xl font-semibold text-white mb-4">Drop evidence files here</h2>
            <form method="post" enctype="multipart/form-data" class="border-2 border-dashed border-gray-600 rounded-lg p-12 text-center">
                <p class="mb-4">or click to browse</p>
                <input type="file" name="file" class="hidden" id="file-input">
                <label for="file-input" class="btn-primary px-6 py-2 rounded-lg cursor-pointer">Select Files</label>
                <button type="submit" class="btn-green px-6 py-2 rounded-lg ml-4">Upload</button>
            </form>
        </div>
        <div class="card p-6 rounded-lg"><h2 class="text-xl font-semibold text-white mb-4">Quick Actions</h2><div class="flex space-x-4"><a href="{{ url_for('manual_carving') }}" class="btn-secondary px-6 py-3 rounded-lg w-full text-center">Manual Carving</a><a href="{{ url_for('auto_carving_setup') }}" class="btn-green px-6 py-3 rounded-lg w-full text-center">Start Auto Carving</a></div></div>
    </div>
    <div class="space-y-8">
        <div class="card p-6 rounded-lg">
            <h2 class="text-xl font-semibold text-white mb-4">Uploaded Files</h2>
            {% if uploaded_files %}
                <div class="space-y-2">
                {% for filename, details in uploaded_files.items() %}
                    <div class="p-2 border-b border-gray-700">
                        <div class="flex justify-between items-center">
                            <div>
                                <span>{{ filename }} ({{ details.size_mb }} MB)</span>
                                {% if details.encryption_status.decrypted_path %}<span class="decryption-badge ml-2">DECRYPTED</span>
                                {% elif details.encryption_status.decrypting %}<span class="decrypting-badge ml-2">DECRYPTING...</span>
                                {% elif details.encryption_status.encrypted %}<span class="encryption-badge ml-2">ENCRYPTED</span>
                                {% endif %}
                            </div>
                            <div class="flex space-x-2">
                                <a href="{{ url_for('remove_file', filename=filename) }}" class="bg-red-600 text-white px-3 py-1 text-xs rounded-lg">Remove</a>
                            </div>
                        </div>
                        {% if details.encryption_status.encrypted and not details.encryption_status.decrypted_path %}
                        <div class="text-xs text-gray-400 mt-2 pl-2">
                            Encryption: {{ details.encryption_status.encryption_type }} - {{ details.encryption_status.description }}
                            {% if details.encryption_status.decrypting %}<br><span class="text-yellow-400">Status: {{ decryption_status.message }}</span>{% endif %}
                        </div>
                        {% endif %}
                    </div>
                {% endfor %}
                </div>
            {% else %}<p class="text-gray-500">No files uploaded yet</p>{% endif %}
        </div>
    </div>
</div>
"""
DELETED_FILES_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Deleted File Recovery (Filesystem Scan)</h1>
<p class="text-gray-400 mb-8">Scans the evidence file's filesystem metadata to find entries for deleted files.</p>
<div class="card p-6 rounded-lg">
    <div class="flex justify-between items-center mb-4">
        <h2 class="text-xl font-semibold text-white">Scan Status</h2>
        <a href="{{ url_for('deleted_files', rescan='true') }}" id="start-scan-btn" class="btn-primary px-4 py-2 rounded-lg text-sm">Start / Rescan</a>
    </div>
    <div id="status-container" class="text-gray-400">
        <p>Status: <span id="scan-status-text" class="font-semibold text-yellow-400">Idle</span></p>
        <p>Message: <span id="scan-message-text"></span></p>
        <p>Files Found: <span id="files-found-count">0</span></p>
    </div>
</div>
<div class="card p-6 rounded-lg mt-8">
    <form action="{{ url_for('download_zip') }}" method="post">
        <input type="hidden" name="file_type" value="deleted">
        <div class="flex justify-between items-center mb-4">
            <h2 class="text-xl font-semibold text-white">Found Deleted Files</h2>
            <button type="submit" class="btn-primary px-4 py-2 rounded-lg text-sm">Download Selected as ZIP</button>
        </div>
        <div class="overflow-x-auto">
            <table class="w-full text-left" id="deleted-files-table">
                <thead>
                    <tr class="border-b border-gray-700">
                        <th class="p-2 w-10"><input type="checkbox" id="select-all-deleted" class="h-4 w-4 rounded bg-gray-700 border-gray-600"></th>
                        <th class="p-2">Filename</th>
                        <th class="p-2">Size (Bytes)</th>
                        <th class="p-2">Inode</th>
                        <th class="p-2">Partition Offset</th>
                        <th class="p-2">Actions</th>
                    </tr>
                </thead>
                <tbody></tbody>
            </table>
            <p id="no-files-message" class="text-gray-500 mt-8 text-center">No deleted files found or scan not performed.</p>
        </div>
    </form>
</div>
<script>
    document.getElementById('select-all-deleted').addEventListener('change', function(e) {
        document.querySelectorAll('.deleted-file-checkbox').forEach(function(checkbox) {
            checkbox.checked = e.target.checked;
        });
    });

    function updateStatus() {
        fetch('/deleted_scan_status')
            .then(response => response.json())
            .then(data => {
                const statusText = document.getElementById('scan-status-text');
                const messageText = document.getElementById('scan-message-text');
                const filesFoundCount = document.getElementById('files-found-count');
                const tableBody = document.querySelector("#deleted-files-table tbody");
                const noFilesMessage = document.getElementById('no-files-message');

                messageText.textContent = data.message;
                statusText.textContent = data.in_progress ? "Scanning..." : (data.complete ? "Scan Complete" : "Idle");
                filesFoundCount.textContent = data.files_found;
                tableBody.innerHTML = '';
                
                if (Object.keys(data.files).length > 0) {
                    noFilesMessage.style.display = 'none';
                    for (const inode in data.files) {
                        const file = data.files[inode];
                        const partitionOffsetHex = `0x${file.fs_offset.toString(16).toUpperCase()}`;
                        const row = `<tr class="border-b border-gray-700 hover:bg-gray-800">
                                <td class="p-2"><input type="checkbox" name="selected_files" value="${file.inode}" class="deleted-file-checkbox h-4 w-4 rounded bg-gray-700 border-gray-600"></td>
                                <td class="p-2 break-all">${file.name}</td>
                                <td class="p-2">${file.size}</td>
                                <td class="p-2">${file.inode}</td>
                                <td class="p-2 font-mono">${partitionOffsetHex}</td>
                                <td class="p-2 flex space-x-2">
                                    <a href="/view_deleted_file/${file.inode}" target="_blank" class="btn-secondary px-3 py-1 text-xs rounded-lg">View</a>
                                    <a href="/hex_view_deleted/${file.inode}" target="_blank" class="btn-green px-3 py-1 text-xs rounded-lg">Hex View</a>
                                    <a href="/download_deleted_file/${file.inode}" class="btn-primary px-3 py-1 text-xs rounded-lg">Download</a>
                                </td>
                            </tr>`;
                        tableBody.insertAdjacentHTML('beforeend', row);
                    }
                } else {
                    if (data.complete) noFilesMessage.textContent = 'Scan complete. No deleted file entries were found.';
                    noFilesMessage.style.display = 'block';
                }

                if (data.in_progress) setTimeout(updateStatus, 2000);
            });
    }
    document.addEventListener('DOMContentLoaded', updateStatus);
</script>
"""

HEX_VIEW_CARVED_FILE_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Hex View: <span class="font-mono text-blue-400">{{ filename }}</span></h1>
<p class="text-gray-400 mb-8">Displaying the full hexadecimal and ASCII representation of the carved file.</p>
<div class="card p-6 rounded-lg mt-8">
    <div class="hex-view p-4 rounded-lg text-sm overflow-auto h-[70vh]">
        <pre>{{ hex_content }}</pre>
    </div>
</div>
"""
AUTO_CARVING_SETUP_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Automatic File Carving Setup</h1>
<p class="text-gray-400 mb-8">Select file types for automated carving based on file signatures.</p>
<form method="post" action="{{ url_for('auto_carving_setup') }}">
<div class="grid grid-cols-1 lg:grid-cols-3 gap-8">
    <div class="lg:col-span-2">
        <div class="card p-6 rounded-lg">
            <div class="flex justify-between items-center mb-4 pb-4 border-b border-gray-700">
                <h2 class="text-xl font-semibold text-white">Select File Types</h2>
                <span id="selected-count" class="bg-blue-600 text-white px-3 py-1 rounded-full text-sm font-medium">0 selected</span>
            </div>
            {% for category, types in signatures.items() %}
            <h3 class="text-lg font-semibold {{ colors[loop.index0 % colors|length] }} mt-6 mb-2 border-b border-gray-700 pb-1">{{ category }}</h3>
            <div class="grid grid-cols-2 md:grid-cols-3 gap-4">
                {% for name, sig in types.items() %}
                <div class="flex items-center">
                    <input type="checkbox" name="file_types" value="{{ name }}" id="{{ name }}" class="file-type-checkbox h-4 w-4 rounded bg-gray-700 border-gray-600 text-blue-600 focus:ring-blue-500">
                    <label for="{{ name }}" class="ml-3 text-white cursor-pointer">{{ name }}</label>
                </div>
                {% endfor %}
            </div>
            {% endfor %}
        </div>
    </div>
    <div class="lg:col-span-1">
        <div class="card p-6 rounded-lg sticky top-8">
             <h3 class="text-lg font-semibold text-white mb-4">Begin Carving</h3>
            <p class="text-gray-400 text-sm mb-4">Once you have selected the desired file types, start the carving process.</p>
            <button type="submit" class="btn-green w-full py-3 rounded-lg font-semibold">Start Auto Carving</button>
        </div>
    </div>
</div>
</form>
<script>
document.addEventListener('DOMContentLoaded', function() {
    const checkboxes = document.querySelectorAll('.file-type-checkbox');
    const counter = document.getElementById('selected-count');
    function updateCounter() {
        const selectedCount = document.querySelectorAll('.file-type-checkbox:checked').length;
        counter.textContent = `${selectedCount} selected`;
    }
    checkboxes.forEach(checkbox => checkbox.addEventListener('change', updateCounter));
    updateCounter();
});
</script>
"""

AUTO_CARVING_PROCESS_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-2">Automatic File Carving in Progress</h1>
<p class="text-gray-400 mb-8">Scanning evidence file for selected file types.</p>

<div class="card p-6 rounded-lg">
    <div class="flex justify-between items-center mb-4">
        <h2 class="text-xl font-semibold text-white">Carving Progress</h2>
        <span id="progress-percent" class="font-bold text-blue-400">0% Complete</span>
    </div>
    <div class="w-full bg-gray-700 rounded-full h-4">
        <div id="progress-bar" class="bg-blue-600 h-4 rounded-full transition-all duration-500" style="width: 0%"></div>
    </div>
    <p class="text-sm text-gray-400 mt-2">Processing offset: <span id="current-offset">0x00000000</span> | Found: <span id="files-found">0</span> files</p>
</div>

<div class="card p-6 rounded-lg mt-8">
    <h2 class="text-xl font-semibold text-white mb-4">Live Hex Preview of Found Files</h2>
    <div id="live-hex-view" class="hex-view p-4 rounded-lg text-sm overflow-y-auto max-h-[40rem]">
        <p class="text-gray-500">Waiting for files to be found...</p>
    </div>
</div>

<div class="mt-8 flex justify-between items-center">
    <h3 class="text-xl font-semibold text-white">Found Files (<span id="total-found-count">0</span>)</h3>
    <div class="flex space-x-4">
        <a href="{{ url_for('auto_carving_setup') }}" class="btn-secondary px-6 py-2 rounded-lg">Back to Setup</a>
        <a href="{{ url_for('recovered_files') }}" id="view-recovered-btn" class="btn-primary px-6 py-2 rounded-lg opacity-50 pointer-events-none">View Recovered Files</a>
    </div>
</div>

<script>
function updateProgress() {
    fetch('/carving_status')
        .then(response => response.json())
        .then(data => {
            // Update progress bar and text elements
            document.getElementById('progress-bar').style.width = data.progress + '%';
            const progressPercent = document.getElementById('progress-percent');
            progressPercent.innerText = data.progress + '% Complete';
            document.getElementById('current-offset').innerText = data.current_offset;
            document.getElementById('files-found').innerText = data.files_found;
            document.getElementById('total-found-count').innerText = data.files_found;

            // Update the live hex preview with the 5 most recent files
            const liveHexView = document.getElementById('live-hex-view');
            let fullHexContent = '';
            const recentFiles = data.found_files_list.slice(-5).reverse();
            if (recentFiles.length > 0) {
                recentFiles.forEach(file => {
                    fullHexContent += `<p class="text-xs text-green-400">${file.name} @ ${file.offset}</p><pre>${file.hex_preview}</pre><hr class="border-gray-600 my-2">`;
                });
                liveHexView.innerHTML = fullHexContent;
            } else {
                liveHexView.innerHTML = '<p class="text-gray-500">Waiting for files to be found...</p>';
            }

            // Check if the carving process is complete
            if (data.complete) {
                progressPercent.className = 'font-bold text-green-400';
                progressPercent.innerText = 'Scan Complete';
                const viewBtn = document.getElementById('view-recovered-btn');
                viewBtn.classList.remove('opacity-50', 'pointer-events-none');
            } else {
                // If not complete, schedule the next update
                setTimeout(updateProgress, 1000);
            }
        });
}
// Start the update loop when the page loads
document.addEventListener('DOMContentLoaded', updateProgress);
</script>
"""
RECOVERED_FILES_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Recovered Carved Files</h1>
<div class="card p-6 rounded-lg">
    <form action="{{ url_for('download_zip') }}" method="post">
        <input type="hidden" name="file_type" value="carved">
        <div class="flex justify-between items-center mb-4">
            <h2 class="text-xl font-semibold text-white">Recovered Files ({{ carved_files|length }} found)</h2>
            {% if carved_files %}
            <button type="submit" class="btn-primary px-4 py-2 rounded-lg text-sm">Download Selected as ZIP</button>
            {% endif %}
        </div>
        {% if carved_files %}
            <table class="w-full text-left">
                <thead>
                    <tr class="border-b border-gray-700">
                        <th class="p-2 w-10"><input type="checkbox" id="select-all-carved" class="h-4 w-4 rounded bg-gray-700 border-gray-600"></th>
                        <th class="p-2">ID</th>
                        <th class="p-2">Filename</th>
                        <th class="p-2">Offset</th>
                        <th class="p-2">Actions</th>
                    </tr>
                </thead>
                <tbody>
                {% for file in carved_files.values()|sort(attribute='id') %}
                    <tr class="border-b border-gray-700 hover:bg-gray-800">
                        <td class="p-2"><input type="checkbox" name="selected_files" value="{{ file.name }}" class="h-4 w-4 rounded bg-gray-700 border-gray-600 file-checkbox"></td>
                        <td class="p-2">{{ file.id }}</td>
                        <td class="p-2 font-mono">{{ file.name }}</td>
                        <td class="p-2 font-mono">{{ file.offset }}</td>
                        <td class="p-2 flex space-x-2">
                            <a href="{{ url_for('view_carved_file', filename=file.name) }}" target="_blank" class="btn-secondary px-3 py-1 text-xs rounded-lg">View</a>
                            <a href="{{ url_for('hex_view_carved', filename=file.name) }}" target="_blank" class="btn-green px-3 py-1 text-xs rounded-lg">Hex View</a>
                            <a href="{{ url_for('download_carved_file', filename=file.name) }}" class="btn-primary px-3 py-1 text-xs rounded-lg">Download</a>
                        </td>
                    </tr>
                {% endfor %}
                </tbody>
            </table>
        {% else %}
            <p class="text-gray-500">No files recovered yet.</p>
        {% endif %}
    </form>
</div>
<script>
document.getElementById('select-all-carved').addEventListener('change', function(e) {
    document.querySelectorAll('.file-checkbox').forEach(function(checkbox) {
        checkbox.checked = e.target.checked;
    });
});
</script>
"""

MANUAL_CARVING_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Manual Carving</h1>
<p class="text-gray-400 mb-8">View the evidence file and extract data by offset, or find data blocks using headers and footers.</p>
<div class="card p-6 rounded-lg">
    <div class="grid grid-cols-1 md:grid-cols-3 gap-6">
        <div>
            <h2 class="text-xl font-semibold text-white mb-4">1. Carving Tool</h2>
            <form action="{{ url_for('perform_manual_carve') }}" method="POST" class="space-y-4">
                <div>
                    <label for="start_offset" class="block text-sm font-medium text-gray-300">Start Offset (bytes)</label>
                    <input type="text" name="start_offset" id="start_offset" value="0" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
                </div>
                <div>
                    <label for="length" class="block text-sm font-medium text-gray-300">Length (bytes)</label>
                    <input type="text" name="length" id="length" value="4096" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
                </div>
                <button type="submit" class="btn-green w-full px-4 py-2 rounded-lg h-10 font-semibold">Carve Selection</button>
            </form>
        </div>
        <div>
            <h2 class="text-xl font-semibold text-white mb-4">2. Find Value</h2>
            <div class="space-y-4">
                 <div>
                    <label for="search-term" class="block text-sm font-medium text-gray-300">Find (Text or Hex)</label>
                    <input type="text" id="search-term" placeholder="e.g., JFIF or FF D8" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
                </div>
                <div>
                    <label for="search-type" class="block text-sm font-medium text-gray-300">Search As</label>
                    <select id="search-type" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2 h-10">
                        <option value="text">Text</option>
                        <option value="hex">Hex</option>
                    </select>
                </div>
                <div class="flex space-x-4">
                    <button id="find-btn" class="btn-primary w-full px-4 py-2 rounded-lg h-10 font-semibold">Find</button>
                    <button id="find-next-btn" class="btn-secondary w-full px-4 py-2 rounded-lg h-10 font-semibold">Find Next</button>
                </div>
            </div>
        </div>
        <div>
            <h2 class="text-xl font-semibold text-white mb-4">3. Find Data Block</h2>
            <div class="space-y-4">
                 <div>
                    <label for="header-term" class="block text-sm font-medium text-gray-300">Header (Text or Hex)</label>
                    <input type="text" id="header-term" placeholder="e.g., FF D8 FF E0" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
                </div>
                <div>
                    <label for="footer-term" class="block text-sm font-medium text-gray-300">Footer (Text or Hex)</label>
                    <input type="text" id="footer-term" placeholder="e.g., FF D9" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2">
                </div>
                <button id="find-block-btn" class="btn-primary w-full px-4 py-2 rounded-lg h-10 font-semibold">Find Block</button>
            </div>
        </div>
    </div>
</div>

<div class="card p-6 rounded-lg mt-8">
    <h2 class="text-xl font-semibold text-white mb-4">Hex Viewer (Scrollable)</h2>
    <div id="hex-viewer-container" class="hex-view p-4 rounded-lg text-sm overflow-auto h-[60vh]">
        <pre id="hex-viewer-content"></pre>
        <div id="loader" class="text-center p-4 text-gray-400">Loading...</div>
    </div>
</div>

<script>
document.addEventListener('DOMContentLoaded', function() {
    const viewerContainer = document.getElementById('hex-viewer-container');
    const viewerContent = document.getElementById('hex-viewer-content');
    const loader = document.getElementById('loader');
    const startOffsetInput = document.getElementById('start_offset');
    const lengthInput = document.getElementById('length');

    let nextOffset = 0;
    let isLoading = false;
    let lastFoundOffset = -1;

    function loadChunk(offset, clearContent = false) {
        if (isLoading) return;
        isLoading = true;
        loader.style.display = 'block';
        if (clearContent) {
            viewerContent.innerHTML = '';
            viewerContainer.scrollTop = 0;
        }

        fetch(`/get_hex_chunk?offset=${offset}`)
            .then(response => response.json())
            .then(data => {
                if (data.hex_view) {
                    viewerContent.innerHTML += data.hex_view;
                    nextOffset = data.next_offset;
                } else {
                    viewerContainer.removeEventListener('scroll', handleScroll);
                    loader.innerText = 'End of File';
                }
                isLoading = false;
                loader.style.display = 'none';
            });
    }

    function handleScroll() {
        const nearBottom = viewerContainer.scrollTop + viewerContainer.clientHeight >= viewerContainer.scrollHeight - 200;
        if (nearBottom && !isLoading) {
            loadChunk(nextOffset);
        }
    }
    viewerContainer.addEventListener('scroll', handleScroll);

    function performSearch() {
        const term = document.getElementById('search-term').value;
        const type = document.getElementById('search-type').value;
        if (!term) { alert('Please enter a search term.'); return; }
        const startFrom = lastFoundOffset === -1 ? 0 : lastFoundOffset + 1;
        
        const formData = new FormData();
        formData.append('term', term);
        formData.append('type', type);
        formData.append('start_offset', startFrom);

        loader.innerText = "Searching...";
        loader.style.display = 'block';

        fetch('/find_in_file', { method: 'POST', body: formData })
            .then(response => response.json())
            .then(data => {
                if (data.offset >= 0) {
                    lastFoundOffset = data.offset;
                    startOffsetInput.value = data.offset;
                    loadChunk(data.offset, true);
                } else {
                    alert('Search term not found.');
                    lastFoundOffset = -1;
                    loader.style.display = 'none';
                }
            });
    }
    
    function performBlockSearch() {
        const header = document.getElementById('header-term').value;
        const footer = document.getElementById('footer-term').value;
        const type = document.getElementById('search-type').value; // You might want a separate type dropdown for this
        if (!header || !footer) { alert('Please enter both a header and footer.'); return; }
        
        const formData = new FormData();
        formData.append('header_term', header);
        formData.append('footer_term', footer);
        formData.append('type', type);
        
        loader.innerText = "Searching for block...";
        loader.style.display = 'block';

        fetch('/find_block', { method: 'POST', body: formData })
            .then(response => response.json())
            .then(data => {
                if (data.status === 'success') {
                    startOffsetInput.value = data.start_offset;
                    lengthInput.value = data.length;
                    loadChunk(data.start_offset, true);
                } else {
                    alert('Could not find a data block with that header and footer.');
                    loader.style.display = 'none';
                }
            });
    }

    document.getElementById('find-btn').addEventListener('click', () => { lastFoundOffset = -1; performSearch(); });
    document.getElementById('find-next-btn').addEventListener('click', performSearch);
    document.getElementById('find-block-btn').addEventListener('click', performBlockSearch);
    
    loadChunk(0);
});
</script>
"""

FORENSIC_ANALYSIS_CONTENT = """
<h1 class="text-3xl font-bold text-white mb-4">Forensic Analysis</h1>
<p class="text-gray-400 mb-8">High-level analysis of the uploaded evidence file.</p>
<div class="card p-6 rounded-lg">
{% if not uploaded_files %}
    <p class="text-gray-500">No file uploaded. Please upload an evidence file to begin analysis.</p>
{% else %}
    {% for filename, details in uploaded_files.items() %}
    <h2 class="text-2xl font-semibold text-white mb-6">Analysis for: <span class="font-mono">{{ filename }}</span></h2>
    <div class="space-y-8">

        <div>
            <h3 class="text-lg font-semibold text-white mb-2">File Status</h3>
            <div class="text-gray-300">
                {% if details.encryption_status.encrypted %}
                    {% if details.encryption_status.decrypted_path %}
                        <p><span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-sm font-medium bg-green-600 text-green-100">DECRYPTED</span> - Analysis is being performed on the decrypted version of the file.</p>
                    {% else %}
                        <p><span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-sm font-medium bg-red-600 text-red-100">ENCRYPTED</span> - <span class="text-yellow-400">Warning:</span> Analysis results may be inaccurate. Please decrypt the file on the Upload page first.</p>
                        <p class="text-sm text-gray-400 mt-1">Detected Type: {{ details.encryption_status.description }}</p>
                    {% endif %}
                {% else %}
                    <p><span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-sm font-medium bg-gray-600 text-gray-100">NOT ENCRYPTED</span> - The file does not appear to be encrypted.</p>
                {% endif %}
            </div>
        </div>
        <div>
            <h3 class="text-lg font-semibold text-white mb-2">General Information</h3>
            <ul class="list-disc list-inside text-gray-300 space-y-1">
            {% for result in details.forensic_results %}
                <li>{{ result }}</li>
            {% endfor %}
            </ul>
        </div>

        <div>
            <h3 class="text-lg font-semibold text-white mb-2">File Hashes (First 64KB)</h3>
            <ul class="list-disc list-inside text-gray-300 space-y-1 font-mono text-sm">
            {% for algo, hash_value in details.hash_info.items() %}
                <li><strong>{{ algo }}:</strong> {{ hash_value }}</li>
            {% endfor %}
            </ul>
        </div>
        
    </div>
    {% endfor %}
{% endif %}
</div>
"""

VIEW_FILE_CONTENT = """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>Viewing: {{ filename if is_carved else file_info.name }}</title>
<style>body, html { margin: 0; padding: 0; height: 100%; overflow: hidden; background-color: #374151; } embed, iframe, img { width: 100%; height: 100%; border: none; } pre { margin: 0; background-color: #111827; color: #d1d5db; height: 100vh; overflow: auto; padding: 20px; box-sizing: border-box; font-family: monospace; white-space: pre-wrap; word-wrap: break-word; }</style>
</head><body>
    {% if mime_type.startswith('image/') %}<img src="{{ url_for('serve_file_data', type='carved', filename=filename) if is_carved else url_for('serve_file_data', type='deleted', filename=inode) }}" alt="Preview">
    {% elif mime_type == 'application/pdf' %}<iframe src="{{ url_for('serve_file_data', type='carved', filename=filename) if is_carved else url_for('serve_file_data', type='deleted', filename=inode) }}"></iframe>
    {% elif mime_type.startswith('text/') %}<pre>{{ text_content }}</pre>
    {% else %}<div style="padding: 40px; text-align: center; color: white; font-family: sans-serif;"><h1 style="font-size: 2em;">Preview Not Available</h1><p>Direct preview for '{{ mime_type }}' is not supported.</p><a href="{{ url_for('download_carved_file', filename=filename) if is_carved else url_for('download_deleted_file', inode=inode) }}" style="display: inline-block; margin-top: 20px; padding: 10px 20px; background-color: #3b82f6; color: white; text-decoration: none; border-radius: 5px;">Download File</a></div>
    {% endif %}
</body></html>
"""

LOG_VIEWER_TEMPLATE = """
<h1 class="text-3xl font-bold text-white mb-4">{{ title }}</h1>
<p class="text-gray-400 mb-8">{{ description }}</p>

<div class="card p-6 rounded-lg">
    <div class="mb-6 p-4 bg-gray-800 border border-yellow-500 rounded-lg text-sm text-gray-300">
        <p><strong>Note:</strong> This tool will analyze the currently loaded evidence file: <strong class="font-mono">{{ evidence_filename }}</strong>.</p>
        <p class="mt-1">It attempts to parse the entire file as the selected log type. For finding and extracting specific log files from a large disk image, please use the <a href="{{ url_for('auto_carving_setup') }}" class="text-blue-400 hover:underline">Auto Carving</a> feature.</p>
    </div>
    
    <form method="post" class="space-y-4">
        <div>
            <label for="os_select" class="block text-sm font-medium text-gray-300">1. Select Operating System</label>
            <select name="os_name" id="os_select" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2 h-10">
                <option value="">-- Select OS --</option>
                {% for os in os_options.keys() %}
                <option value="{{ os }}">{{ os }}</option>
                {% endfor %}
            </select>
        </div>

        <div>
            <label for="ext_select" class="block text-sm font-medium text-gray-300">2. Select Log Type to Parse As</label>
            <select name="log_ext" id="ext_select" class="mt-1 block w-full bg-gray-800 border-gray-600 rounded-md shadow-sm p-2 h-10" disabled>
                <option>-- Select OS First --</option>
            </select>
        </div>
        
        <button type="submit" class="btn-green w-full px-4 py-2 rounded-lg h-10 font-semibold">Parse Evidence File</button>
    </form>
</div>

{% if log_content %}
<div class="card p-6 rounded-lg mt-8">
    <h2 class="text-xl font-semibold text-white mb-4">Parsing Result for: <span class="font-mono">{{ evidence_filename }}</span></h2>
    <div class="log-view p-4 rounded-lg text-sm overflow-auto h-[60vh]">
        <pre>{{ log_content }}</pre>
    </div>
</div>
{% endif %}

<script>
    const osOptions = {{ os_options|tojson }};
    const osSelect = document.getElementById('os_select');
    const extSelect = document.getElementById('ext_select');

    osSelect.addEventListener('change', function() {
        const selectedOs = this.value;
        extSelect.innerHTML = '';
        if (selectedOs && osOptions[selectedOs]) {
            extSelect.disabled = false;
            osOptions[selectedOs].forEach(function(ext) {
                const option = new Option(ext || 'No Extension', ext);
                extSelect.add(option);
            });
        } else {
            extSelect.disabled = true;
            extSelect.add(new Option('-- Select OS First --', ''));
        }
    });
</script>
"""

# --- Flask Routes ---

@app.route('/')
def index():
    return redirect(url_for('evidence_upload'))

@app.route('/evidence_upload', methods=['GET', 'POST'])
def evidence_upload():
    global uploaded_files_db
    modal_payload = {"show_encryption_modal": False}
    if request.method == 'POST' and 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        uploaded_files_db.clear(); carved_files_db.clear(); deleted_files_db.clear()
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        encryption_info = detect_encryption(filepath)
        uploaded_files_db[filename] = {
            "path": filepath, "size_mb": f"{os.path.getsize(filepath) / (1024*1024):.2f}",
            "encryption_status": {
                "encrypted": encryption_info.get('encrypted'), "encryption_type": encryption_info.get('encryption_type'),
                "description": encryption_info.get('description'), "decrypting": False, "decrypted_path": None,
            },
            "forensic_results": perform_forensic_analysis(filepath), "hash_info": show_hash_info(filepath)
        }
        flash(f"Successfully uploaded and analyzed '{filename}'.")
        
        modal_payload['show_encryption_modal'] = True
        if encryption_info.get('encrypted'):
            modal_payload['modal_title'] = "File Encrypted"
            modal_payload['modal_body'] = f"""
            <p>This file is encrypted. Decrypt it before analysis.</p>
            <form action='{url_for("decrypt_file", filename=filename)}' method='POST' class='mt-4'>
                <input type='password' name='password' placeholder='Enter Password (if known)' class='w-full p-2 rounded bg-gray-800 border border-gray-600 text-white'>
                <div class='flex justify-end space-x-4 mt-4'>
                    <button type='submit' name='action' value='no_password' class='btn-secondary px-4 py-2 rounded-lg'>Decrypt (Dictionary)</button>
                    <button type='submit' name='action' value='with_password' class='btn-primary px-4 py-2 rounded-lg'>Decrypt with Password</button>
                </div>
            </form>"""
        else:
            modal_payload['modal_title'] = "File Status"
            modal_payload['modal_body'] = f"""
            <p>This file is not encrypted. You can proceed to the analysis tools.</p>
            <div class='flex justify-end space-x-4 mt-4'>
                 <button onclick="document.getElementById('encryptionModal').style.display='none'" class='btn-green px-4 py-2 rounded-lg'>Continue</button>
            </div>
            """
            
    content = render_template_string(EVIDENCE_UPLOAD_CONTENT, uploaded_files=uploaded_files_db, decryption_status=decryption_status)
    return render_template_string(BASE_TEMPLATE, content=content, **modal_payload)

@app.route('/reporting', methods=['GET', 'POST'])
def reporting():
    """Handles the generation of a complete forensic report in multiple formats."""
    if not uploaded_files_db:
        flash("Please upload an evidence file first to generate a report.", "error")
        return redirect(url_for('evidence_upload'))

    if request.method == 'POST':
        # Gather common data for the report
        case_details = {
            "name": request.form.get('case_name', 'N/A'),
            "number": request.form.get('case_number', 'N/A'),
            "examiner": request.form.get('examiner_name', 'N/A')
        }
        report_format = request.form.get('report_format')
        now = datetime.datetime.now()
        safe_case_name = secure_filename(case_details['name'])

        # --- Logic for different report formats ---
        if report_format == 'html':
            for filename, info in carved_files_db.items():
                filepath = os.path.join(app.config['CARVED_FOLDER'], filename)
                info['thumbnail_uri'] = create_thumbnail_data_uri(filepath)
            
            report_html = render_template_string(
                REPORT_TEMPLATE, case_details=case_details, evidence_file=uploaded_files_db,
                carved_files=carved_files_db, deleted_files=deleted_files_db, now=now
            )
            response = make_response(report_html)
            response.headers['Content-Disposition'] = f"attachment; filename=Report_{safe_case_name}.html"
            return response

        elif report_format == 'pdf':
            if not HTML:
                flash("PDF generation requires the 'weasyprint' library. Please run: pip install weasyprint", "error")
                return redirect(url_for('reporting'))
            
            for filename, info in carved_files_db.items():
                filepath = os.path.join(app.config['CARVED_FOLDER'], filename)
                info['thumbnail_uri'] = create_thumbnail_data_uri(filepath)

            report_html = render_template_string(
                REPORT_TEMPLATE, case_details=case_details, evidence_file=uploaded_files_db,
                carved_files=carved_files_db, deleted_files=deleted_files_db, now=now
            )
            pdf_bytes = HTML(string=report_html).write_pdf()
            response = make_response(pdf_bytes)
            response.headers['Content-Disposition'] = f"attachment; filename=Report_{safe_case_name}.pdf"
            response.headers['Content-Type'] = 'application/pdf'
            return response
            
        elif report_format == 'docx':
            docx_buffer = generate_docx_report_data(case_details, uploaded_files_db, carved_files_db, deleted_files_db, now)
            if not docx_buffer:
                flash("DOCX generation requires the 'python-docx' library. Please run: pip install python-docx", "error")
                return redirect(url_for('reporting'))
            return send_file(docx_buffer, download_name=f'Report_{safe_case_name}.docx', as_attachment=True)

        elif report_format == 'csv':
            zip_buffer = generate_csv_zip_report_data(case_details, uploaded_files_db, carved_files_db, deleted_files_db)
            return send_file(zip_buffer, download_name=f'Report_{safe_case_name}_CSV.zip', as_attachment=True, mimetype='application/zip')

    # For a GET request, show the page with the form
    content = render_template_string(REPORTING_PAGE_CONTENT, now=datetime.datetime.now())
    return render_template_string(BASE_TEMPLATE, content=content)

def generate_csv_zip_report_data(case_details, evidence_file, carved_files, deleted_files):
    """Generate a ZIP file containing CSV reports for case details, evidence, carved files, and deleted files."""
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Case details CSV
        case_csv = io.StringIO()
        writer = csv.writer(case_csv)
        writer.writerow(['Case Name', 'Case Number', 'Examiner'])
        writer.writerow([case_details.get('name', ''), case_details.get('number', ''), case_details.get('examiner', '')])
        zf.writestr('case_details.csv', case_csv.getvalue())

        # Evidence file CSV
        evidence_csv = io.StringIO()
        writer = csv.writer(evidence_csv)
        writer.writerow(['Filename', 'Size (MB)', 'MD5', 'SHA-1', 'SHA-256'])
        for filename, details in evidence_file.items():
            writer.writerow([
                filename,
                details.get('size_mb', ''),
                details.get('hash_info', {}).get('MD5', ''),
                details.get('hash_info', {}).get('SHA-1', ''),
                details.get('hash_info', {}).get('SHA-256', ''),
            ])
        zf.writestr('evidence_files.csv', evidence_csv.getvalue())

        # Carved files CSV
        carved_csv = io.StringIO()
        writer = csv.writer(carved_csv)
        writer.writerow(['ID', 'Filename', 'Offset', 'Size (KB)'])
        for file in sorted(carved_files.values(), key=lambda x: x.get('id', 0)):
            writer.writerow([
                file.get('id', ''),
                file.get('name', ''),
                file.get('offset', ''),
                file.get('size_kb', ''),
            ])
        zf.writestr('carved_files.csv', carved_csv.getvalue())

        # Deleted files CSV
        deleted_csv = io.StringIO()
        writer = csv.writer(deleted_csv)
        writer.writerow(['Inode', 'Filename', 'Size (Bytes)', 'Modified', 'Accessed', 'Created'])
        for file in deleted_files.values():
            writer.writerow([
                file.get('inode', ''),
                file.get('name', ''),
                file.get('size', ''),
                file.get('mtime', ''),
                file.get('atime', ''),
                file.get('ctime', ''),
            ])
        zf.writestr('deleted_files.csv', deleted_csv.getvalue())

    memory_file.seek(0)
    return memory_file

@app.route('/hex_view_deleted/<int:inode>')
def hex_view_deleted(inode):
    """Displays a hex view of a single deleted file recovered via pytsk3."""
    if not uploaded_files_db or inode not in deleted_files_db:
        flash("File not found or session expired. Please start a new session.", "error")
        return redirect(url_for('deleted_files'))

    try:
        # Get file metadata from our database and the evidence path
        file_info = deleted_files_db[inode]
        evidence_path = list(uploaded_files_db.values())[0]['path']
        
        # Use pytsk3 to read the file content directly from the disk image
        img = pytsk3.Img_Info(evidence_path)
        fs = pytsk3.FS_Info(img, offset=file_info['fs_offset'])
        fs_file = fs.open_meta(inode=inode)
        content = fs_file.read_random(0, file_info['size'])
        
        # Format the binary content into a readable hex string
        hex_content = format_hex_view(content)
        
        # Reuse the existing hex view template for a consistent look
        content_html = render_template_string(HEX_VIEW_CARVED_FILE_CONTENT, filename=file_info['name'], hex_content=hex_content)
        return render_template_string(BASE_TEMPLATE, content=content_html)

    except Exception as e:
        flash(f"An error occurred while reading inode {inode}: {e}", "error")
        return redirect(url_for('deleted_files'))
    
@app.route('/download_zip', methods=['POST'])
def download_zip():
    """Handles the creation and download of a ZIP archive of selected files."""
    if not uploaded_files_db:
        flash("No evidence file is loaded. Please start a new session.", "error")
        return redirect(url_for('evidence_upload'))

    selected_ids = request.form.getlist('selected_files')
    file_type = request.form.get('file_type')

    if not selected_ids:
        flash("No files were selected for download.", "warning")
        if file_type == 'carved':
            return redirect(url_for('recovered_files'))
        else:
            return redirect(url_for('deleted_files'))

    memory_file = io.BytesIO()
    zip_filename = f"carved_files_{datetime.datetime.now():%Y%m%d%H%M%S}.zip"

    try:
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            if file_type == 'carved':
                for filename in selected_ids:
                    s_filename = secure_filename(filename)
                    filepath = os.path.join(app.config['CARVED_FOLDER'], s_filename)
                    if os.path.exists(filepath):
                        zf.write(filepath, arcname=s_filename)
            
            elif file_type == 'deleted':
                zip_filename = f"deleted_files_{datetime.datetime.now():%Y%m%d%H%M%S}.zip"
                evidence_path = list(uploaded_files_db.values())[0]['path']
                img = pytsk3.Img_Info(evidence_path)
                
                for inode_str in selected_ids:
                    inode = int(inode_str)
                    if inode in deleted_files_db:
                        file_info = deleted_files_db[inode]
                        try:
                            fs = pytsk3.FS_Info(img, offset=file_info['fs_offset'])
                            fs_file = fs.open_meta(inode=inode)
                            content = fs_file.read_random(0, file_info['size'])
                            
                            # Sanitize filename for the zip archive
                            safe_arcname = secure_filename(f"inode_{inode}_{file_info['name']}")
                            if not safe_arcname:
                                safe_arcname = f"inode_{inode}_unnamed_file"

                            zf.writestr(safe_arcname, content)
                        except Exception as e:
                            print(f"Could not recover inode {inode}: {e}")
        
        memory_file.seek(0)
        return send_file(memory_file, download_name=zip_filename, as_attachment=True, mimetype='application/zip')

    except Exception as e:
        flash(f"An error occurred while creating the ZIP file: {e}", "error")
        if file_type == 'carved':
            return redirect(url_for('recovered_files'))
        else:
            return redirect(url_for('deleted_files'))
        
@app.route('/hex_view_carved/<filename>')
def hex_view_carved(filename):
    """Displays a hex view of a single carved file."""
    if not uploaded_files_db:
        flash("No evidence file is loaded. Please start a new session.", "error")
        return redirect(url_for('evidence_upload'))

    s_filename = secure_filename(filename)
    filepath = os.path.join(app.config['CARVED_FOLDER'], s_filename)

    if not os.path.exists(filepath):
        flash(f"Carved file '{s_filename}' not found.", "error")
        return redirect(url_for('recovered_files'))

    try:
        with open(filepath, 'rb') as f:
            content = f.read()
        
        hex_content = format_hex_view(content)
        
        content_html = render_template_string(HEX_VIEW_CARVED_FILE_CONTENT, filename=s_filename, hex_content=hex_content)
        return render_template_string(BASE_TEMPLATE, content=content_html)

    except Exception as e:
        flash(f"Error reading or processing file '{s_filename}': {e}", "error")
        return redirect(url_for('recovered_files'))

@app.route('/remove_file/<filename>')
def remove_file(filename):
    if filename in uploaded_files_db:
        del uploaded_files_db[filename]
        flash(f'File {filename} removed from session.', 'success')
    return redirect(url_for('evidence_upload'))

@app.route('/decrypt_file/<filename>', methods=['POST'])
def decrypt_file(filename):
    if filename not in uploaded_files_db:
        flash("File not found.", 'error')
        return redirect(url_for('evidence_upload'))
    password = request.form.get('password')
    action = request.form.get('action')
    file_info = uploaded_files_db[filename]
    if action == 'with_password' and not password:
        flash("Password field was empty.", 'error')
        return redirect(url_for('evidence_upload'))
    threading.Thread(target=attempt_decryption, args=(file_info['path'], file_info['encryption_status']['encryption_type'], None if action == 'no_password' else password)).start()
    flash("Decryption process started. The page will update automatically.", 'info')
    return redirect(url_for('evidence_upload'))

@app.route('/deleted_files')
def deleted_files():
    rescan = request.args.get('rescan', 'false').lower() == 'true'
    if not uploaded_files_db:
        flash("Please upload an evidence file first.")
        return redirect(url_for('evidence_upload'))
    if rescan and not deleted_scan_status.get("in_progress"):
        filepath = list(uploaded_files_db.values())[0]['path']
        threading.Thread(target=scan_for_deleted_files_engine, args=(filepath,)).start()
    content = render_template_string(DELETED_FILES_CONTENT)
    return render_template_string(BASE_TEMPLATE, content=content)

@app.route('/deleted_scan_status')
def deleted_scan_status_endpoint():
    return jsonify({"in_progress": deleted_scan_status["in_progress"], "complete": deleted_scan_status["complete"], "files_found": deleted_scan_status["files_found"], "message": deleted_scan_status["message"], "files": deleted_files_db})

@app.route('/view_deleted_file/<int:inode>')
def view_deleted_file(inode):
    if not uploaded_files_db or inode not in deleted_files_db: return "File not found", 404
    file_info = deleted_files_db[inode]
    mime_type, _ = mimetypes.guess_type(file_info['name'])
    mime_type = mime_type or 'application/octet-stream'
    
    if mime_type.startswith('text/'):
        try:
            filepath = list(uploaded_files_db.values())[0]['path']
            img = pytsk3.Img_Info(filepath)
            fs = pytsk3.FS_Info(img, offset=file_info['fs_offset'])
            fs_file = fs.open_meta(inode=inode)
            content = fs_file.read_random(0, min(file_info['size'], 10000))
            text_content = content.decode('utf-8', errors='ignore')
            return render_template_string(VIEW_FILE_CONTENT, is_carved=False, file_info=file_info, inode=inode, mime_type=mime_type, text_content=text_content)
        except Exception as e: return f"Error reading text file: {e}", 500
    
    return render_template_string(VIEW_FILE_CONTENT, is_carved=False, file_info=file_info, inode=inode, mime_type=mime_type)

@app.route('/download_deleted_file/<int:inode>')
def download_deleted_file(inode):
    if not uploaded_files_db or inode not in deleted_files_db: return "File not found", 404
    file_info = deleted_files_db[inode]
    filepath = list(uploaded_files_db.values())[0]['path']
    try:
        img = pytsk3.Img_Info(filepath)
        fs = pytsk3.FS_Info(img, offset=file_info['fs_offset'])
        fs_file = fs.open_meta(inode=inode)
        content = fs_file.read_random(0, file_info['size'])
        response = make_response(content)
        response.headers.set('Content-Type', 'application/octet-stream')
        response.headers.set('Content-Disposition', 'attachment', filename=f"recovered_{inode}_{file_info['name']}")
        return response
    except Exception as e: return f"Error recovering file: {e}", 500

@app.route('/auto_carving_setup', methods=['GET', 'POST'])
def auto_carving_setup():
    if request.method == 'POST':
        selected_types = request.form.getlist('file_types')
        if not uploaded_files_db: flash("Please upload an evidence file first."); return redirect(url_for('evidence_upload'))
        if not selected_types: flash("Please select at least one file type to carve."); return redirect(url_for('auto_carving_setup'))
        evidence_file_path = list(uploaded_files_db.values())[0]['path']
        threading.Thread(target=simple_file_carver, args=(evidence_file_path, selected_types)).start()
        return redirect(url_for('auto_carving_process'))
    # Add this list of Tailwind CSS color classes
    colors = ['text-sky-400', 'text-emerald-400', 'text-amber-400', 'text-rose-400', 'text-violet-400', 'text-teal-400', 'text-cyan-400', 'text-lime-400', 'text-pink-400', 'text-indigo-400']
# Pass the colors list into the template
    content = render_template_string(AUTO_CARVING_SETUP_CONTENT, signatures=FILE_SIGNATURES, colors=colors)
    return render_template_string(BASE_TEMPLATE, content=content)

@app.route('/auto_carving_process')
def auto_carving_process():
    content = render_template_string(AUTO_CARVING_PROCESS_CONTENT)
    return render_template_string(BASE_TEMPLATE, content=content)

@app.route('/carving_status')
def carving_status_endpoint():
    return jsonify(carving_status)

@app.route('/recovered_files')
def recovered_files():
    content = render_template_string(RECOVERED_FILES_CONTENT, carved_files=carved_files_db)
    return render_template_string(BASE_TEMPLATE, content=content)

@app.route('/view_carved_file/<filename>')
def view_carved_file(filename):
    filepath = os.path.join(app.config['CARVED_FOLDER'], secure_filename(filename))
    if not os.path.exists(filepath): return "File not found", 404
    mime_type, _ = mimetypes.guess_type(filepath)
    mime_type = mime_type or 'application/octet-stream'
    text_content = ""
    if mime_type.startswith('text/'):
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read(10000)
        except Exception as e: text_content = f"Could not read file as text: {e}"
    return render_template_string(VIEW_FILE_CONTENT, is_carved=True, filename=filename, mime_type=mime_type, text_content=text_content)

@app.route('/serve_file_data/<type>/<path:filename>')
def serve_file_data(type, filename):
    if type == 'carved':
        return send_from_directory(app.config['CARVED_FOLDER'], secure_filename(filename))
    elif type == 'deleted':
        inode = int(filename)
        if not uploaded_files_db or inode not in deleted_files_db: return "File not found", 404
        file_info = deleted_files_db[inode]
        filepath = list(uploaded_files_db.values())[0]['path']
        try:
            img = pytsk3.Img_Info(filepath)
            fs = pytsk3.FS_Info(img, offset=file_info['fs_offset'])
            fs_file = fs.open_meta(inode=inode)
            content = fs_file.read_random(0, file_info['size'])
            mime_type, _ = mimetypes.guess_type(file_info['name'])
            mime_type = mime_type or 'application/octet-stream'
            return Response(content, mimetype=mime_type)
        except Exception as e: return f"Error serving file data: {e}", 500
    return "Invalid file type", 400

@app.route('/download_carved_file/<filename>')
def download_carved_file(filename):
    return send_from_directory(app.config['CARVED_FOLDER'], secure_filename(filename), as_attachment=True)

@app.route('/manual_carving')
def manual_carving():
    if not uploaded_files_db:
        flash("Please upload an evidence file first to use the manual carver.", "error")
        return redirect(url_for('evidence_upload'))
    content = render_template_string(MANUAL_CARVING_CONTENT)
    return render_template_string(BASE_TEMPLATE, content=content)

@app.route('/get_hex_chunk')
def get_hex_chunk():
    if not uploaded_files_db:
        return jsonify({"error": "No file uploaded"})
    
    CHUNK_SIZE = 4096
    offset = request.args.get('offset', 0, type=int)
    filepath = list(uploaded_files_db.values())[0]['path']
    file_size = os.path.getsize(filepath)
    
    if offset >= file_size:
        return jsonify({"hex_view": "", "next_offset": offset})

    with open(filepath, 'rb') as f:
        f.seek(offset)
        data_chunk = f.read(CHUNK_SIZE)
        hex_view = format_hex_view(data_chunk, start_offset=offset)
    
    return jsonify({"hex_view": hex_view, "next_offset": offset + len(data_chunk)})

@app.route('/find_in_file', methods=['POST'])
def find_in_file():
    if not uploaded_files_db:
        return jsonify({"offset": -1, "error": "No file uploaded"})
    
    term = request.form.get('term')
    search_type = request.form.get('type')
    start_offset = request.form.get('start_offset', 0, type=int)

    try:
        search_bytes = bytes.fromhex(term.replace(" ", "")) if search_type == 'hex' else term.encode()
    except ValueError:
        return jsonify({"offset": -1, "error": "Invalid Hex sequence."})

    filepath = list(uploaded_files_db.values())[0]['path']
    
    try:
        with open(filepath, 'rb') as f:
            with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                found_pos = mm.find(search_bytes, start_offset)
                return jsonify({"offset": found_pos})
    except Exception as e:
        return jsonify({"offset": -1, "error": str(e)})

@app.route('/find_block', methods=['POST'])
def find_block():
    if not uploaded_files_db:
        return jsonify({"status": "error", "message": "No file uploaded"})
    
    header_term = request.form.get('header_term')
    footer_term = request.form.get('footer_term')
    search_type = request.form.get('type')

    try:
        header_bytes = bytes.fromhex(header_term.replace(" ", "")) if search_type == 'hex' else header_term.encode()
        footer_bytes = bytes.fromhex(footer_term.replace(" ", "")) if search_type == 'hex' else footer_term.encode()
    except ValueError:
        return jsonify({"status": "error", "message": "Invalid Hex sequence."})

    filepath = list(uploaded_files_db.values())[0]['path']
    
    try:
        with open(filepath, 'rb') as f:
            with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                header_pos = mm.find(header_bytes, 0)
                if header_pos != -1:
                    footer_pos = mm.find(footer_bytes, header_pos + len(header_bytes))
                    if footer_pos != -1:
                        block_len = (footer_pos + len(footer_bytes)) - header_pos
                        return jsonify({"status": "success", "start_offset": header_pos, "length": block_len})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

    return jsonify({"status": "not_found"})

@app.route('/perform_manual_carve', methods=['POST'])
def perform_manual_carve():
    if not uploaded_files_db:
        flash("No evidence file is loaded.", "error")
        return redirect(url_for('manual_carving'))

    try:
        start_offset = int(request.form.get('start_offset', 0))
        length = int(request.form.get('length', 0))
    except (ValueError, TypeError):
        flash("Invalid offset or length. Please enter numbers only.", "error")
        return redirect(url_for('manual_carving'))

    if length <= 0:
        flash("Length must be a positive number.", "error")
        return redirect(url_for('manual_carving', offset=start_offset))

    filepath = list(uploaded_files_db.values())[0]['path']
    file_size = os.path.getsize(filepath)

    if start_offset >= file_size:
        flash("Start offset is beyond the end of the file.", "error")
        return redirect(url_for('manual_carving', offset=start_offset))

    with open(filepath, 'rb') as f:
        f.seek(start_offset)
        data_to_carve = f.read(min(length, file_size - start_offset))

    carved_filename = f"manual_carve_{start_offset}_{length}.dat"
    save_path = os.path.join(app.config['CARVED_FOLDER'], carved_filename)
    with open(save_path, 'wb') as out_f:
        out_f.write(data_to_carve)

    file_id = len(carved_files_db) + len(deleted_files_db) + 1
    carved_files_db[carved_filename] = {"id": file_id, "name": carved_filename, "offset": f"0x{start_offset:08X}", "size_kb": f"{len(data_to_carve)/1024:.2f} KB"}

    flash(f"Successfully carved {len(data_to_carve)} bytes to '{carved_filename}'.", "success")
    return redirect(url_for('recovered_files'))

@app.route('/forensic_analysis')
def forensic_analysis():
    content = render_template_string(FORENSIC_ANALYSIS_CONTENT, uploaded_files=uploaded_files_db)
    return render_template_string(BASE_TEMPLATE, content=content)

# --- NEW ROUTES FOR LOG VIEWERS ---
@app.route('/log_file_viewer', methods=['GET', 'POST'])
def log_file_viewer():
    if not uploaded_files_db:
        flash("Please upload an evidence file first to use the log viewer.", "error")
        return redirect(url_for('evidence_upload'))

    evidence_filename = list(uploaded_files_db.keys())[0]
    log_content = None

    if request.method == 'POST':
        filepath = list(uploaded_files_db.values())[0]['path']
        log_ext = request.form.get('log_ext', '').lower()

        try:
            if not log_ext and '.' in evidence_filename:
                log_ext = '.' + evidence_filename.rsplit('.', 1)[1].lower()

            if log_ext in [".log", ".txt"] or log_ext.startswith('/var/log/'):
                log_content = parse_text_log(filepath)
            elif log_ext == ".csv":
                log_content = parse_csv_log(filepath)
            elif log_ext == ".json":
                log_content = parse_json_log(filepath)
            elif log_ext == ".xml":
                log_content = parse_xml_log(filepath)
            elif log_ext in [".evt", ".evtx"]:
                log_content = parse_evtx_log(filepath)
            else:
                log_content = f"[INFO] Format '{log_ext}' requires specialized handling. Displaying raw text.\n\n"
                log_content += parse_text_log(filepath)
        except Exception as e:
            log_content = f"[ERROR] Failed to parse file as '{log_ext}'. It may be part of a larger disk image.\n\nDetails: {e}"

    template_vars = {
        "title": "Log File Viewer",
        "description": "Analyze the loaded evidence file by parsing it as a specific log type.",
        "os_options": LOG_OS_OPTIONS,
        "log_content": log_content,
        "evidence_filename": evidence_filename
    }
    content = render_template_string(LOG_VIEWER_TEMPLATE, **template_vars)
    return render_template_string(BASE_TEMPLATE, content=content)


@app.route('/event_log_viewer', methods=['GET', 'POST'])
def event_log_viewer():
    if not uploaded_files_db:
        flash("Please upload an evidence file first to use the event log viewer.", "error")
        return redirect(url_for('evidence_upload'))

    evidence_filename = list(uploaded_files_db.keys())[0]
    log_content = None
    
    if request.method == 'POST':
        filepath = list(uploaded_files_db.values())[0]['path']
        log_ext = request.form.get('log_ext', '').lower()

        try:
            if not log_ext and '.' in evidence_filename:
                log_ext = '.' + evidence_filename.rsplit('.', 1)[1].lower()

            if log_ext in [".evtx", ".evt"]:
                log_content = parse_evtx_log(filepath)
            elif log_ext.startswith("/var/log/") or log_ext in [".log", ".txt", ".asl"]:
                log_content = parse_text_log(filepath)
            else:
                log_content = f"[INFO] Format '{log_ext}' requires external tools. Displaying as plain text.\n\n"
                log_content += parse_text_log(filepath)
        except Exception as e:
            log_content = f"[ERROR] Failed to parse file as '{log_ext}'. It may be part of a larger disk image.\n\nDetails: {e}"

    template_vars = {
        "title": "Event Log Viewer",
        "description": "Analyze the loaded evidence file by parsing it as a specific event log type.",
        "os_options": EVENT_OS_OPTIONS,
        "log_content": log_content,
        "evidence_filename": evidence_filename
    }
    content = render_template_string(LOG_VIEWER_TEMPLATE, **template_vars)
    return render_template_string(BASE_TEMPLATE, content=content)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)